"""Professional DOCX generation for Career Copilot.

The layout intentionally follows the CV and LOM reference documents supplied
for Week 6.  Content remains data-driven, while typography, spacing, section
order and visual hierarchy are deterministic.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
import re
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
PROFILE_PHOTO = PROJECT_ROOT / "assets" / "profile_photo.jpeg"

DEFAULT_CONTACT = {
    "location": "Gaildorf",
    "phone": "+49 176 5695 0651",
    "email": "arjunv8@gmail.com",
}


# ---------------------------------------------------------------------------
# Low-level formatting helpers
# ---------------------------------------------------------------------------

def _set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_paragraph_spacing(paragraph, *, before=0, after=0, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def _add_bottom_border(paragraph, size="6", space="1"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), "000000")


def _set_run_font(run, name="Times New Roman", size=10.0, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def _add_bullet(document: Document, text: str, *, size=9.2, left=0.18, hanging=0.14):
    p = document.add_paragraph()
    p.paragraph_format.left_indent = Inches(left)
    p.paragraph_format.first_line_indent = Inches(-hanging)
    _set_paragraph_spacing(p, after=0.7, line=1.0)
    bullet = p.add_run("•")
    _set_run_font(bullet, size=size)
    spacer = p.add_run("  ")
    _set_run_font(spacer, size=size)
    run = p.add_run(str(text).strip())
    _set_run_font(run, size=size)
    return p


def _add_section_heading(document: Document, title: str):
    p = document.add_paragraph()
    _set_paragraph_spacing(p, before=3.2, after=2.2, line=1.0)
    run = p.add_run(title.upper())
    _set_run_font(run, size=10.2, bold=True)
    _add_bottom_border(p, size="5", space="1")
    return p


def _clean_items(items):
    output = []
    for item in items or []:
        text = str(item or "").strip()
        if text and text not in output:
            output.append(text)
    return output


def _contact(candidate_data: dict) -> dict:
    return {
        "location": str(candidate_data.get("location") or DEFAULT_CONTACT["location"]).strip(),
        "phone": str(candidate_data.get("phone") or DEFAULT_CONTACT["phone"]).strip(),
        "email": str(candidate_data.get("email") or DEFAULT_CONTACT["email"]).strip(),
    }


def _experience_lookup(cv_data: dict) -> dict[tuple[str, str], dict]:
    lookup = {}
    for item in cv_data.get("experience", []) or []:
        if not isinstance(item, dict):
            continue
        company = str(item.get("company", "")).strip().lower()
        role = str(item.get("job_title", "")).strip().lower()
        if company or role:
            lookup[(company, role)] = item
    return lookup


def _find_tailored_experience(lookup: dict, company: str, role: str):
    key = (company.strip().lower(), role.strip().lower())
    if key in lookup:
        return lookup[key]
    # The AI may normalize punctuation or legal suffixes. Use conservative
    # containment only when one side is otherwise unambiguous.
    for (c, r), value in lookup.items():
        if c and r and (
            (c in key[0] or key[0] in c)
            and (r in key[1] or key[1] in r)
        ):
            return value
    return None


# ---------------------------------------------------------------------------
# CV skill grouping - deterministic structure matching the supplied CV
# ---------------------------------------------------------------------------

_SKILL_GROUPS = [
    (
        "Software Engineering & Architecture -",
        (
            "c++", "python", "pybind11", "embedded c", "autosar", "object-oriented",
            "oop", "cmake", "architecture", "uml", "mbsd", "mbse", "vibe coding",
        ),
    ),
    (
        "System & Safety Engineering-",
        (
            "e/e", "requirements", "doors", "iso 26262", "asil", "aspice", "fmea",
            "fta", "risk", "code review", "safety-critical", "system engineering",
        ),
    ),
    (
        "Full-Stack & AI Tools-",
        (
            "react", "typescript", "fastapi", "sqlite", "prompt", "mvp", "ai ",
        ),
    ),
    (
        "Simulation & Engineering Toolchains",
        (
            "matlab", "simulink", "targetlink", "dspace", "scalexio", "microautobox",
            "carmaker", "canoe", "canape", "enterprise architect", "systemdesk",
            "lauterbach", "renesas", "aurelion", "simulation", "hil", "sil", "dil",
        ),
    ),
    (
        "Process & Infrastructure:",
        (
            "agile", "safe", "ci/cd", "jenkins", "conan", "docker", "artifactory",
            "git", "jira", "v-model", "verification", "validation",
        ),
    ),
]


def _all_candidate_terms(candidate_data: dict, prioritized: list[str]) -> list[str]:
    terms = list(prioritized or [])
    for field in ("technical_skills", "programming_languages", "tools", "methodologies"):
        terms.extend(_clean_items(candidate_data.get(field, [])))
    # Stable de-duplication.
    seen = set()
    result = []
    for term in terms:
        key = re.sub(r"[^a-z0-9+#]+", "", term.lower())
        if key and key not in seen:
            seen.add(key)
            result.append(term)
    return result


def _group_skills(candidate_data: dict, prioritized: list[str]) -> list[tuple[str, list[str]]]:
    terms = _all_candidate_terms(candidate_data, prioritized)
    groups: list[tuple[str, list[str]]] = []
    used = set()
    for label, keywords in _SKILL_GROUPS:
        selected = []
        # Prioritized terms naturally occur earlier in `terms`, so the job-
        # relevant technologies appear first while the reference structure stays stable.
        for term in terms:
            low = f" {term.lower()} "
            if any(keyword in low for keyword in keywords):
                key = term.lower().strip()
                if key not in used:
                    used.add(key)
                    selected.append(term)
        if selected:
            groups.append((label, selected))
    return groups


# ---------------------------------------------------------------------------
# CV generation
# ---------------------------------------------------------------------------

def create_cv_docx(cv_data: dict, candidate_data: dict) -> bytes:
    """Create an application-ready CV following the supplied two-page design."""

    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.38)
    section.left_margin = Inches(0.48)
    section.right_margin = Inches(0.48)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(9.2)

    # Header block: name + photograph, then centered contact details and rule.
    table = document.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(7.35)
    table.columns[1].width = Inches(0.75)
    table.cell(0, 0).width = Inches(7.35)
    table.cell(0, 1).width = Inches(0.75)
    table.cell(0, 0).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    table.cell(0, 1).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for cell in table.rows[0].cells:
        _set_cell_margins(cell, top=0, start=0, bottom=0, end=0)

    name = str(candidate_data.get("full_name") or "ARJUN VISWANATHAN").strip().upper()
    parts = name.split(maxsplit=1)
    name_p = table.cell(0, 0).paragraphs[0]
    _set_paragraph_spacing(name_p, after=0)
    first = name_p.add_run(parts[0] if parts else name)
    _set_run_font(first, name="Arial", size=30, bold=True)
    if len(parts) > 1:
        second = name_p.add_run("  " + parts[1])
        _set_run_font(second, name="Arial", size=30, bold=False)

    if PROFILE_PHOTO.exists():
        photo_p = table.cell(0, 1).paragraphs[0]
        photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = photo_p.add_run()
        run.add_picture(str(PROFILE_PHOTO), width=Inches(0.50), height=Inches(0.90))

    contact = _contact(candidate_data)
    contact_p = document.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(contact_p, before=0, after=1.0)
    contact_run = contact_p.add_run(
        f"{contact['location']} | Phone: {contact['phone']} | E-Mail: {contact['email']}"
    )
    _set_run_font(contact_run, name="Arial", size=8.3, italic=True)
    _add_bottom_border(contact_p, size="12", space="3")

    summary = str(cv_data.get("professional_summary") or candidate_data.get("professional_summary") or "").strip()
    if summary:
        p = document.add_paragraph()
        _set_paragraph_spacing(p, before=6.0, after=3.0, line=1.08)
        r = p.add_run(summary)
        _set_run_font(r, size=9.2)

    _add_section_heading(document, "Technical Skills")
    for label, skills in _group_skills(candidate_data, cv_data.get("prioritized_skills", [])):
        label_p = document.add_paragraph()
        _set_paragraph_spacing(label_p, before=1.2, after=0.2)
        label_run = label_p.add_run(label)
        _set_run_font(label_run, size=9.2, bold=True)
        _add_bullet(document, ", ".join(skills), size=9.0)

    _add_section_heading(document, "Work Experience")
    tailored_lookup = _experience_lookup(cv_data)
    for entry in candidate_data.get("employment_history", []) or []:
        if not isinstance(entry, dict):
            continue
        company = str(entry.get("company", "")).strip()
        role = str(entry.get("role", "")).strip()
        start = str(entry.get("start_date", "")).strip()
        end = str(entry.get("end_date", "")).strip()

        company_p = document.add_paragraph()
        _set_paragraph_spacing(company_p, before=2.0, after=0.3)
        # Keep location in parentheses visually secondary.
        company_match = re.match(r"^(.*?)\s*(\([^)]*\))$", company)
        if company_match:
            r1 = company_p.add_run(company_match.group(1).strip() + " ")
            _set_run_font(r1, size=9.2, bold=True)
            r2 = company_p.add_run(company_match.group(2))
            _set_run_font(r2, size=9.2)
        else:
            r1 = company_p.add_run(company)
            _set_run_font(r1, size=9.2, bold=True)

        role_p = document.add_paragraph()
        _set_paragraph_spacing(role_p, after=0.8)
        role_run = role_p.add_run(role)
        _set_run_font(role_run, size=9.2, bold=True, italic=True)
        if start or end:
            dates = f" ({start}-{end})" if start and end else f" ({start or end})"
            date_run = role_p.add_run(dates)
            _set_run_font(date_run, size=9.2, italic=True)

        tailored = _find_tailored_experience(tailored_lookup, company, role)
        bullets = tailored.get("bullets", []) if tailored else entry.get("responsibilities", [])
        for bullet in _clean_items(bullets):
            _add_bullet(document, bullet, size=8.9, left=0.18, hanging=0.14)

    _add_section_heading(document, "Education")
    for entry in candidate_data.get("education", []) or []:
        if not isinstance(entry, dict):
            continue
        year = str(entry.get("end_date", "")).strip()
        degree = str(entry.get("degree", "")).strip()
        field = str(entry.get("field_of_study", "")).strip()
        institution = str(entry.get("institution", "")).strip()
        p = document.add_paragraph()
        _set_paragraph_spacing(p, after=1.0)
        yr = p.add_run((year + "   ") if year else "")
        _set_run_font(yr, size=9.1, bold=True)
        desc = ", ".join(x for x in (f"{degree} {field}".strip(), institution) if x)
        rr = p.add_run(desc + ("." if desc and not desc.endswith(".") else ""))
        _set_run_font(rr, size=9.1)

    _add_section_heading(document, "Patents & Publications & Awards")
    subsections = [
        ("Patents", candidate_data.get("patents", [])),
        ("Publications", candidate_data.get("publications", [])),
        ("Awards", candidate_data.get("awards", [])),
        ("Certifications", candidate_data.get("certifications", [])),
    ]
    for label, items in subsections:
        items = _clean_items(items)
        if not items:
            continue
        p = document.add_paragraph()
        _set_paragraph_spacing(p, before=1.0, after=0.2)
        r = p.add_run(label)
        _set_run_font(r, size=9.1, bold=True)
        for item in items:
            _add_bullet(document, item, size=8.9)

    languages = _clean_items(candidate_data.get("spoken_languages", []))
    if languages:
        _add_section_heading(document, "Languages")
        for item in languages:
            _add_bullet(document, item, size=8.9)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Cover letter generation
# ---------------------------------------------------------------------------

def _body_paragraphs(text: str) -> list[str]:
    """Extract body paragraphs while removing accidental LLM wrapper lines."""
    raw_parts = [part.strip() for part in re.split(r"\n\s*\n", str(text or "")) if part.strip()]
    cleaned = []
    for part in raw_parts:
        low = part.lower().strip()
        if low.startswith(("to,", "to:", "subject:", "dear ", "sincerely", "best regards", "kind regards")):
            continue
        cleaned.append(re.sub(r"\s*\n\s*", " ", part).strip())
    if cleaned:
        return cleaned
    single = re.sub(r"\s*\n\s*", " ", str(text or "")).strip()
    return [single] if single else []


def create_cover_letter_docx(cover_letter: str, candidate_data: dict, job_data: dict) -> bytes:
    """Create a one-page professional LOM matching the supplied reference."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)

    name = str(candidate_data.get("full_name") or "ARJUN VISWANATHAN").strip().upper()
    parts = name.split(maxsplit=1)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, after=3.5)
    first = p.add_run(parts[0] if parts else name)
    _set_run_font(first, size=17.5, bold=True)
    if len(parts) > 1:
        last = p.add_run("   " + parts[1])
        _set_run_font(last, size=17.5)

    contact = _contact(candidate_data)
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_paragraph_spacing(p, after=24)
    r = p.add_run(f"{contact['email']}  •  {contact['phone']}")
    _set_run_font(r, size=8.5)

    company = str(job_data.get("company") or "Hiring Company").strip()
    job_title = str(job_data.get("job_title") or "the advertised position").strip()
    job_id = str(job_data.get("job_id") or "").strip()

    p = document.add_paragraph()
    _set_paragraph_spacing(p, after=13)
    r = p.add_run("To,")
    _set_run_font(r, size=11)

    p = document.add_paragraph()
    _set_paragraph_spacing(p, after=14)
    r = p.add_run(company)
    _set_run_font(r, size=11, bold=True)

    subject = f"Subject: Application for {job_title}"
    if job_id:
        subject += f" – Job ID {job_id}"
    p = document.add_paragraph()
    _set_paragraph_spacing(p, after=14)
    r = p.add_run(subject)
    _set_run_font(r, size=11, bold=True)

    p = document.add_paragraph()
    _set_paragraph_spacing(p, after=14)
    r = p.add_run("Dear Hiring Team,")
    _set_run_font(r, size=11, bold=True)

    for body in _body_paragraphs(cover_letter):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _set_paragraph_spacing(p, after=11, line=1.05)
        r = p.add_run(body)
        _set_run_font(r, size=10.6)

    p = document.add_paragraph()
    _set_paragraph_spacing(p, before=4, after=1)
    r = p.add_run("Sincerely")
    _set_run_font(r, size=10.8)
    p = document.add_paragraph()
    _set_paragraph_spacing(p, after=0)
    r = p.add_run(str(candidate_data.get("full_name") or "Arjun Viswanathan").title())
    _set_run_font(r, size=10.8, bold=True)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
