import json
import os
from datetime import datetime
from io import BytesIO
from pathlib import Path

import fitz
import pandas as pd
import streamlit as st
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from pydantic import BaseModel

# ------------------------------------------------------------
# OPTIONAL C++ MODULE IMPORT
# ------------------------------------------------------------

try:
    import match_engine

    MATCH_ENGINE_AVAILABLE = True
    MATCH_ENGINE_IMPORT_ERROR = None

except Exception as import_error:
    match_engine = None
    MATCH_ENGINE_AVAILABLE = False
    MATCH_ENGINE_IMPORT_ERROR = str(import_error)


# ------------------------------------------------------------
# CONFIGURATION
# ------------------------------------------------------------

load_dotenv()

OPENAI_MODEL = os.getenv(
    "OPENAI_MODEL",
    "gpt-4.1-mini",
)

DEFAULT_PROFILE_PATH = Path(
    "candidate_profile.json"
)

APPLICATIONS_FOLDER = Path(
    "data/applications"
)


# ------------------------------------------------------------
# PYDANTIC DATA MODELS
# ------------------------------------------------------------

class EmploymentEntry(BaseModel):
    company: str
    role: str
    start_date: str
    end_date: str
    responsibilities: list[str]
    achievements: list[str]
    technologies: list[str]


class EducationEntry(BaseModel):
    institution: str
    degree: str
    field_of_study: str
    start_date: str
    end_date: str


class CandidateProfile(BaseModel):
    full_name: str
    professional_title: str
    professional_summary: str

    technical_skills: list[str]
    programming_languages: list[str]
    tools: list[str]
    methodologies: list[str]
    industries: list[str]

    employment_history: list[EmploymentEntry]
    education: list[EducationEntry]

    certifications: list[str]
    spoken_languages: list[str]
    leadership_experience: list[str]
    achievements: list[str]


class JobProfile(BaseModel):
    job_title: str
    company: str
    location: str

    required_skills: list[str]
    preferred_skills: list[str]

    required_experience_years: float

    education_requirements: list[str]
    language_requirements: list[str]

    responsibilities: list[str]

    industry: str
    seniority: str


class TailoredExperience(BaseModel):
    company: str
    job_title: str
    bullets: list[str]


class TailoredCV(BaseModel):
    professional_title: str
    professional_summary: str
    prioritized_skills: list[str]
    experience: list[TailoredExperience]
    education: list[str]
    certifications: list[str]


class InterviewQuestion(BaseModel):
    question: str
    category: str
    why_it_matters: str
    preparation_points: list[str]


class InterviewPreparation(BaseModel):
    technical_questions: list[InterviewQuestion]
    behavioral_questions: list[InterviewQuestion]
    experience_questions: list[InterviewQuestion]


# ------------------------------------------------------------
# SESSION STATE
# ------------------------------------------------------------

SESSION_DEFAULTS = {
    "candidate_profile": None,
    "candidate_source": None,
    "extracted_text": None,
    "uploaded_cv_name": None,
    "job_profile": None,
    "job_description_text": "",
    "match_result": None,
    "tailored_cv": None,
    "cover_letter": None,
    "interview_preparation": None,
}


def initialize_session_state():
    for key, value in SESSION_DEFAULTS.items():
        if key not in st.session_state:
            st.session_state[key] = value


def clear_application_outputs():
    """
    Clear all information that becomes stale when either the candidate
    profile or job profile changes.
    """

    st.session_state["match_result"] = None
    st.session_state["tailored_cv"] = None
    st.session_state["cover_letter"] = None
    st.session_state["interview_preparation"] = None


def clear_job_and_application_outputs():
    """
    Use this when the candidate changes. A match generated for the old
    candidate must not remain attached to the new candidate.
    """

    st.session_state["job_profile"] = None
    st.session_state["job_description_text"] = ""
    clear_application_outputs()


# ------------------------------------------------------------
# OPENAI CLIENT
# ------------------------------------------------------------

def get_openai_client() -> OpenAI:
    api_key = os.getenv(
        "OPENAI_API_KEY"
    )

    if not api_key:
        raise ValueError(
            "OPENAI_API_KEY was not found in the .env file."
        )

    return OpenAI(
        api_key=api_key
    )


# ------------------------------------------------------------
# CV EXTRACTION
# ------------------------------------------------------------

def extract_pdf_text(
    file_bytes: bytes,
) -> str:

    document = fitz.open(
        stream=file_bytes,
        filetype="pdf",
    )

    try:
        pages = []

        for page in document:
            text = page.get_text(
                "text"
            )

            if text.strip():
                pages.append(
                    text
                )

        return "\n".join(
            pages
        )

    finally:
        document.close()


def extract_docx_text(
    file_bytes: bytes,
) -> str:

    document = Document(
        BytesIO(file_bytes)
    )

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    return "\n".join(
        paragraphs
    )


def extract_cv_text(
    filename: str,
    file_bytes: bytes,
) -> str:

    lower_filename = (
        filename.lower()
    )

    if lower_filename.endswith(
        ".pdf"
    ):
        return extract_pdf_text(
            file_bytes
        )

    if lower_filename.endswith(
        ".docx"
    ):
        return extract_docx_text(
            file_bytes
        )

    raise ValueError(
        "Unsupported file type. "
        "Upload a PDF or DOCX file."
    )


# ------------------------------------------------------------
# EXISTING PROFILE LOAD / SAVE
# ------------------------------------------------------------

def validate_candidate_profile_data(
    data,
) -> dict:
    """
    Validate arbitrary JSON-like data using the same CandidateProfile
    schema used by the AI extractor.

    This makes loading an existing profile safe and guarantees that the
    rest of the application receives the expected structure.
    """

    if not isinstance(
        data,
        dict,
    ):
        raise ValueError(
            "Candidate profile must be a JSON object."
        )

    validated = (
        CandidateProfile.model_validate(
            data
        )
    )

    return validated.model_dump()


def load_candidate_profile_from_bytes(
    file_bytes: bytes,
) -> dict:

    try:
        raw_data = json.loads(
            file_bytes.decode(
                "utf-8"
            )
        )

    except UnicodeDecodeError as error:
        raise ValueError(
            "The candidate profile must be a UTF-8 JSON file."
        ) from error

    except json.JSONDecodeError as error:
        raise ValueError(
            "The selected file is not valid JSON."
        ) from error

    return validate_candidate_profile_data(
        raw_data
    )


def load_default_candidate_profile():
    """
    Load candidate_profile.json from the project folder when it exists.
    """

    if not DEFAULT_PROFILE_PATH.exists():
        raise FileNotFoundError(
            "candidate_profile.json was not found in the project folder."
        )

    with DEFAULT_PROFILE_PATH.open(
        "r",
        encoding="utf-8",
    ) as file:
        data = json.load(
            file
        )

    return validate_candidate_profile_data(
        data
    )


def save_candidate_profile_to_default_file(
    profile_data: dict,
):
    """
    Save the currently approved profile to candidate_profile.json.
    """

    validated = (
        validate_candidate_profile_data(
            profile_data
        )
    )

    with DEFAULT_PROFILE_PATH.open(
        "w",
        encoding="utf-8",
    ) as file:
        json.dump(
            validated,
            file,
            indent=2,
            ensure_ascii=False,
        )


# ------------------------------------------------------------
# AI PROFILE EXTRACTION
# ------------------------------------------------------------

def create_candidate_profile(
    cv_text: str,
) -> CandidateProfile:

    client = get_openai_client()

    system_prompt = """
You are a precise CV information-extraction system.

Extract only information explicitly supported by the supplied CV.

Rules:
1. Never invent experience, skills, qualifications or achievements.
2. Preserve company names, role names and dates as written.
3. When information is unavailable, use an empty string or empty list.
4. Do not evaluate whether the candidate is good or bad.
5. Do not improve or rewrite the candidate's experience.
6. Do not calculate information that cannot be confirmed.
7. Keep responsibilities and achievements separate where possible.
8. Return information according to the supplied structure.
"""

    response = client.responses.parse(
        model=OPENAI_MODEL,
        input=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": (
                    "Extract a candidate profile from this CV:\n\n"
                    + cv_text
                ),
            },
        ],
        text_format=CandidateProfile,
    )

    profile = (
        response.output_parsed
    )

    if profile is None:
        raise ValueError(
            "The AI did not return a valid candidate profile."
        )

    return profile


# ------------------------------------------------------------
# JOB ANALYSIS
# ------------------------------------------------------------

def create_job_profile(
    job_description: str,
) -> JobProfile:

    client = get_openai_client()

    system_prompt = """
You are a job-description information-extraction system.

Extract structured information from the supplied job advertisement.

Rules:
1. Extract only information supported by the job description.
2. Do not invent requirements.
3. Separate required skills from preferred skills.
4. If years of experience are not explicitly stated, use 0.
5. If company, location, industry or seniority cannot be identified,
   use an empty string.
6. Keep skill names concise.
7. Preserve important technical terminology.
8. Return information according to the provided schema.
"""

    response = client.responses.parse(
        model=OPENAI_MODEL,
        input=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": (
                    "Extract a structured job profile "
                    "from the following job advertisement:\n\n"
                    + job_description
                ),
            },
        ],
        text_format=JobProfile,
    )

    job_profile = (
        response.output_parsed
    )

    if job_profile is None:
        raise ValueError(
            "The AI did not return a valid job profile."
        )

    return job_profile


# ------------------------------------------------------------
# MATCHING
# ------------------------------------------------------------

def calculate_cpp_match(
    candidate_data: dict,
    job_data: dict,
) -> dict:

    if not MATCH_ENGINE_AVAILABLE:
        raise RuntimeError(
            "The C++ match_engine module could not be imported. "
            f"Import error: {MATCH_ENGINE_IMPORT_ERROR}"
        )

    candidate_skills = (
        candidate_data.get(
            "technical_skills",
            [],
        )
    )

    required_skills = (
        job_data.get(
            "required_skills",
            [],
        )
    )

    result = (
        match_engine.calculate_match(
            candidate_skills,
            required_skills,
        )
    )

    return {
        "score": float(
            result.score
        ),
        "matched_skills": list(
            result.matched_skills
        ),
        "missing_skills": list(
            result.missing_skills
        ),
    }


# ------------------------------------------------------------
# APPLICATION GENERATION
# ------------------------------------------------------------

def create_tailored_cv(
    candidate_data: dict,
    job_data: dict,
    match_data: dict,
) -> TailoredCV:

    client = get_openai_client()

    system_prompt = """
You are the CV tailoring component of Career Copilot.

Create a job-specific CV using ONLY facts contained in the approved
candidate profile.

You may:
- prioritize relevant experience
- reorder skills according to job relevance
- rewrite existing experience more clearly
- shorten less relevant content
- emphasize evidence that supports the job requirements

You must NEVER:
- invent skills
- invent employers
- invent projects
- invent responsibilities
- invent certifications
- invent achievements
- claim experience merely because the job requires it

The approved candidate profile is the factual source of truth.

The compatibility analysis is guidance only.
If a skill appears in the missing-skills list, do not claim that the
candidate has that skill.

Return the tailored CV according to the provided schema.
"""

    user_content = f"""
APPROVED CANDIDATE PROFILE:

{json.dumps(candidate_data, indent=2)}

STRUCTURED JOB PROFILE:

{json.dumps(job_data, indent=2)}

COMPATIBILITY ANALYSIS:

{json.dumps(match_data, indent=2)}

Create a tailored CV for this job.
"""

    response = client.responses.parse(
        model=OPENAI_MODEL,
        input=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_content,
            },
        ],
        text_format=TailoredCV,
    )

    tailored_cv = (
        response.output_parsed
    )

    if tailored_cv is None:
        raise ValueError(
            "The AI did not return a valid tailored CV."
        )

    return tailored_cv


def create_cover_letter(
    candidate_data: dict,
    job_data: dict,
    match_data: dict,
) -> str:

    client = get_openai_client()

    system_prompt = """
You are the cover-letter generation component of Career Copilot.

Create a concise and professional cover letter for the supplied job.

Rules:
1. Use only facts supported by the approved candidate profile.
2. Never invent experience, skills, achievements, certifications
   or responsibilities.
3. Focus on the strongest genuine overlap between the candidate
   and the job requirements.
4. Use the compatibility analysis to identify matched strengths.
5. Do not claim missing skills.
6. Avoid generic wording where possible.
7. Keep the letter approximately 300-450 words.
8. Make the tone professional and confident.
"""

    user_content = f"""
APPROVED CANDIDATE PROFILE:

{json.dumps(candidate_data, indent=2)}

JOB PROFILE:

{json.dumps(job_data, indent=2)}

COMPATIBILITY ANALYSIS:

{json.dumps(match_data, indent=2)}

Write the cover letter.
"""

    response = client.responses.create(
        model=OPENAI_MODEL,
        input=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_content,
            },
        ],
    )

    cover_letter = (
        response.output_text
    )

    if not cover_letter:
        raise ValueError(
            "No cover letter was returned."
        )

    return cover_letter


def create_interview_preparation(
    candidate_data: dict,
    job_data: dict,
    match_data: dict,
) -> InterviewPreparation:

    client = get_openai_client()

    system_prompt = """
You are the interview-preparation component of Career Copilot.

Generate interview questions specifically for the candidate and job.

Create:
- 5 technical questions
- 5 experience-based questions
- 5 behavioral questions

For every question provide:
- the question
- its category
- why the interviewer may ask it
- preparation points

Rules:
1. Base questions on actual job requirements.
2. Use candidate experience where relevant.
3. Pay special attention to missing skills.
4. Do not invent experience.
5. If a missing skill is important, explain that the candidate should
   prepare the theory rather than falsely claim experience.
6. Keep preparation advice concise and practical.
"""

    user_content = f"""
CANDIDATE PROFILE:

{json.dumps(candidate_data, indent=2)}

JOB PROFILE:

{json.dumps(job_data, indent=2)}

COMPATIBILITY ANALYSIS:

{json.dumps(match_data, indent=2)}

Generate the interview-preparation package.
"""

    response = client.responses.parse(
        model=OPENAI_MODEL,
        input=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_content,
            },
        ],
        text_format=InterviewPreparation,
    )

    interview_prep = (
        response.output_parsed
    )

    if interview_prep is None:
        raise ValueError(
            "The AI did not return valid interview preparation."
        )

    return interview_prep


# ------------------------------------------------------------
# DOCX GENERATION
# ------------------------------------------------------------

def create_cv_docx(
    cv_data: dict,
) -> bytes:

    document = Document()

    document.add_heading(
        cv_data.get(
            "professional_title",
            "Tailored CV",
        ),
        level=0,
    )

    document.add_heading(
        "Professional Summary",
        level=1,
    )

    document.add_paragraph(
        cv_data.get(
            "professional_summary",
            "",
        )
    )

    document.add_heading(
        "Key Skills",
        level=1,
    )

    for skill in cv_data.get(
        "prioritized_skills",
        [],
    ):
        document.add_paragraph(
            str(skill),
            style="List Bullet",
        )

    document.add_heading(
        "Professional Experience",
        level=1,
    )

    for experience in cv_data.get(
        "experience",
        [],
    ):

        if not isinstance(
            experience,
            dict,
        ):
            continue

        document.add_heading(
            (
                experience.get(
                    "job_title",
                    "",
                )
                + " — "
                + experience.get(
                    "company",
                    "",
                )
            ),
            level=2,
        )

        for bullet in experience.get(
            "bullets",
            [],
        ):
            document.add_paragraph(
                str(bullet),
                style="List Bullet",
            )

    document.add_heading(
        "Education",
        level=1,
    )

    for item in cv_data.get(
        "education",
        [],
    ):
        document.add_paragraph(
            str(item),
            style="List Bullet",
        )

    document.add_heading(
        "Certifications",
        level=1,
    )

    for item in cv_data.get(
        "certifications",
        [],
    ):
        document.add_paragraph(
            str(item),
            style="List Bullet",
        )

    buffer = BytesIO()
    document.save(
        buffer
    )
    buffer.seek(0)

    return buffer.getvalue()


def create_cover_letter_docx(
    cover_letter: str,
) -> bytes:

    document = Document()

    document.add_heading(
        "Cover Letter",
        level=0,
    )

    for paragraph in cover_letter.split(
        "\n"
    ):

        if paragraph.strip():
            document.add_paragraph(
                paragraph.strip()
            )

    buffer = BytesIO()
    document.save(
        buffer
    )
    buffer.seek(0)

    return buffer.getvalue()


def create_interview_docx(
    prep: dict,
) -> bytes:

    if not isinstance(
        prep,
        dict,
    ):
        raise TypeError(
            "Interview preparation must be structured data."
        )

    document = Document()

    document.add_heading(
        "Interview Preparation",
        level=0,
    )

    sections = [
        (
            "Technical Questions",
            "technical_questions",
        ),
        (
            "Experience Questions",
            "experience_questions",
        ),
        (
            "Behavioral Questions",
            "behavioral_questions",
        ),
    ]

    for section_title, key in sections:

        document.add_heading(
            section_title,
            level=1,
        )

        for item in prep.get(
            key,
            [],
        ):

            if not isinstance(
                item,
                dict,
            ):
                continue

            document.add_heading(
                item.get(
                    "question",
                    "Question",
                ),
                level=2,
            )

            document.add_paragraph(
                "Why this may be asked:"
            )

            document.add_paragraph(
                item.get(
                    "why_it_matters",
                    "",
                )
            )

            document.add_paragraph(
                "Preparation points:"
            )

            for point in item.get(
                "preparation_points",
                [],
            ):
                document.add_paragraph(
                    str(point),
                    style="List Bullet",
                )

    buffer = BytesIO()
    document.save(
        buffer
    )
    buffer.seek(0)

    return buffer.getvalue()


# ------------------------------------------------------------
# APPLICATION HISTORY
# ------------------------------------------------------------

def save_application_record(
    candidate_data: dict,
    job_data: dict,
    match_data: dict,
    tailored_cv=None,
    cover_letter=None,
    interview_prep=None,
    status="Prepared",
):

    APPLICATIONS_FOLDER.mkdir(
        parents=True,
        exist_ok=True,
    )

    timestamp = datetime.now()

    application_id = (
        timestamp.strftime(
            "%Y%m%d_%H%M%S"
        )
    )

    application_record = {
        "application_id": application_id,
        "created_at": timestamp.isoformat(),
        "company": job_data.get(
            "company",
            "",
        ),
        "job_title": job_data.get(
            "job_title",
            "",
        ),
        "location": job_data.get(
            "location",
            "",
        ),
        "match_score": match_data.get(
            "score",
            0,
        ),
        "matched_skills": match_data.get(
            "matched_skills",
            [],
        ),
        "missing_skills": match_data.get(
            "missing_skills",
            [],
        ),
        "status": status,
        "candidate_profile": candidate_data,
        "job_profile": job_data,
        "tailored_cv": tailored_cv,
        "cover_letter": cover_letter,
        "interview_preparation": interview_prep,
    }

    file_path = (
        APPLICATIONS_FOLDER
        / f"{application_id}.json"
    )

    with file_path.open(
        "w",
        encoding="utf-8",
    ) as file:
        json.dump(
            application_record,
            file,
            indent=2,
            ensure_ascii=False,
        )

    return file_path


def load_application_history():

    if not APPLICATIONS_FOLDER.exists():
        return []

    applications = []

    for file_path in (
        APPLICATIONS_FOLDER.glob(
            "*.json"
        )
    ):

        try:
            with file_path.open(
                "r",
                encoding="utf-8",
            ) as file:
                data = json.load(
                    file
                )

            applications.append(
                data
            )

        except Exception:
            # A single damaged history record should not crash the UI.
            continue

    applications.sort(
        key=lambda item: item.get(
            "created_at",
            "",
        ),
        reverse=True,
    )

    return applications


# ------------------------------------------------------------
# UI HELPERS
# ------------------------------------------------------------

def show_error(
    user_message: str,
    error: Exception,
):

    st.error(
        user_message
    )

    with st.expander(
        "Technical details"
    ):
        st.code(
            str(error)
        )


def candidate_is_ready() -> bool:
    return isinstance(
        st.session_state.get(
            "candidate_profile"
        ),
        dict,
    )


def job_is_ready() -> bool:
    return isinstance(
        st.session_state.get(
            "job_profile"
        ),
        dict,
    )


def match_is_ready() -> bool:
    return isinstance(
        st.session_state.get(
            "match_result"
        ),
        dict,
    )


def render_interview_questions(
    prep_data: dict,
):

    sections = [
        (
            "Technical Questions",
            "technical_questions",
        ),
        (
            "Experience Questions",
            "experience_questions",
        ),
        (
            "Behavioral Questions",
            "behavioral_questions",
        ),
    ]

    for title, key in sections:

        st.markdown(
            f"### {title}"
        )

        questions = prep_data.get(
            key,
            [],
        )

        if not questions:
            st.info(
                f"No {title.lower()} were generated."
            )
            continue

        for index, item in enumerate(
            questions,
            start=1,
        ):

            if not isinstance(
                item,
                dict,
            ):
                continue

            question = item.get(
                "question",
                "Question",
            )

            with st.expander(
                f"{index}. {question}"
            ):

                st.markdown(
                    "**Why it may be asked**"
                )

                st.write(
                    item.get(
                        "why_it_matters",
                        "",
                    )
                )

                st.markdown(
                    "**Preparation points**"
                )

                for point in item.get(
                    "preparation_points",
                    [],
                ):
                    st.write(
                        f"- {point}"
                    )


# ------------------------------------------------------------
# STREAMLIT APPLICATION
# ------------------------------------------------------------

st.set_page_config(
    page_title="Career Copilot",
    page_icon="💼",
    layout="wide",
)

initialize_session_state()

st.title(
    "Career Copilot"
)

st.caption(
    "Beta v0.4"
)

st.warning(
    "AI-generated application content must be reviewed before use."
)

profile_tab, job_tab, application_tab, history_tab = st.tabs(
    [
        "👤 Candidate Profile",
        "🔎 Job Analysis",
        "📄 Application Package",
        "📚 Application History",
    ]
)


# ============================================================
# CANDIDATE PROFILE TAB
# ============================================================

with profile_tab:

    st.header(
        "Candidate Profile"
    )

    st.write(
        "Use an existing approved Career Copilot profile, "
        "or analyse a PDF/DOCX CV to create a new one."
    )

    candidate_mode = st.radio(
        "Choose how to provide the candidate profile",
        [
            "Load existing profile",
            "Analyse CV",
        ],
        horizontal=True,
        key="candidate_input_mode",
    )

    # --------------------------------------------------------
    # OPTION A: LOAD EXISTING PROFILE
    # --------------------------------------------------------

    if (
        candidate_mode
        == "Load existing profile"
    ):

        st.subheader(
            "Load Existing Candidate Profile"
        )

        st.caption(
            "Use a candidate_profile.json file previously exported "
            "by Career Copilot."
        )

        existing_profile_file = st.file_uploader(
            "Choose candidate profile JSON",
            type=["json"],
            key="existing_candidate_profile_upload",
        )

        col_load_file, col_load_default = st.columns(
            2
        )

        with col_load_file:

            load_uploaded_disabled = (
                existing_profile_file
                is None
            )

            if st.button(
                "Load Selected Profile",
                type="primary",
                disabled=load_uploaded_disabled,
                use_container_width=True,
            ):

                try:
                    loaded_profile = (
                        load_candidate_profile_from_bytes(
                            existing_profile_file.getvalue()
                        )
                    )

                    st.session_state[
                        "candidate_profile"
                    ] = loaded_profile

                    st.session_state[
                        "candidate_source"
                    ] = (
                        f"Loaded JSON: "
                        f"{existing_profile_file.name}"
                    )

                    clear_job_and_application_outputs()

                    st.success(
                        "Existing candidate profile loaded successfully."
                    )

                    st.rerun()

                except Exception as error:
                    show_error(
                        "The candidate profile could not be loaded.",
                        error,
                    )

        with col_load_default:

            if st.button(
                "Load candidate_profile.json",
                disabled=(
                    not DEFAULT_PROFILE_PATH.exists()
                ),
                use_container_width=True,
            ):

                try:
                    loaded_profile = (
                        load_default_candidate_profile()
                    )

                    st.session_state[
                        "candidate_profile"
                    ] = loaded_profile

                    st.session_state[
                        "candidate_source"
                    ] = (
                        "Loaded from project candidate_profile.json"
                    )

                    clear_job_and_application_outputs()

                    st.success(
                        "Project candidate profile loaded successfully."
                    )

                    st.rerun()

                except Exception as error:
                    show_error(
                        "The project candidate profile could not be loaded.",
                        error,
                    )

        if not DEFAULT_PROFILE_PATH.exists():
            st.caption(
                "No candidate_profile.json currently exists "
                "in the project folder."
            )

    # --------------------------------------------------------
    # OPTION B: ANALYSE CV
    # --------------------------------------------------------

    else:

        st.subheader(
            "Analyse CV"
        )

        uploaded_file = st.file_uploader(
            "Upload your CV",
            type=[
                "pdf",
                "docx",
            ],
            key="cv_upload",
        )

        if uploaded_file is not None:

            file_identity = (
                uploaded_file.name,
                len(
                    uploaded_file.getvalue()
                ),
            )

            previous_identity = (
                st.session_state.get(
                    "uploaded_cv_identity"
                )
            )

            if (
                previous_identity
                != file_identity
            ):

                try:
                    extracted_text = (
                        extract_cv_text(
                            filename=uploaded_file.name,
                            file_bytes=uploaded_file.getvalue(),
                        )
                    )

                    if not extracted_text.strip():
                        st.error(
                            "No readable text was found in this document."
                        )

                    else:
                        st.session_state[
                            "extracted_text"
                        ] = extracted_text

                        st.session_state[
                            "uploaded_cv_name"
                        ] = uploaded_file.name

                        st.session_state[
                            "uploaded_cv_identity"
                        ] = file_identity

                except Exception as error:
                    show_error(
                        "The CV could not be processed.",
                        error,
                    )

            extracted_text = (
                st.session_state.get(
                    "extracted_text"
                )
            )

            if extracted_text:

                st.success(
                    "The CV text was extracted successfully."
                )

                with st.expander(
                    "View extracted CV text"
                ):
                    st.text_area(
                        "Extracted text",
                        value=extracted_text,
                        height=450,
                        disabled=True,
                    )

                if st.button(
                    "Create Candidate Profile",
                    type="primary",
                ):

                    with st.spinner(
                        "The AI is analysing the CV..."
                    ):

                        try:
                            profile = (
                                create_candidate_profile(
                                    extracted_text
                                )
                            )

                            st.session_state[
                                "candidate_profile"
                            ] = (
                                profile.model_dump()
                            )

                            st.session_state[
                                "candidate_source"
                            ] = (
                                f"AI extracted from "
                                f"{uploaded_file.name}"
                            )

                            clear_job_and_application_outputs()

                            st.success(
                                "Candidate profile created successfully."
                            )

                            st.rerun()

                        except Exception as error:
                            show_error(
                                "Candidate profile generation failed.",
                                error,
                            )

    # --------------------------------------------------------
    # ACTIVE CANDIDATE PROFILE
    # --------------------------------------------------------

    st.divider()

    profile_data = st.session_state.get(
        "candidate_profile"
    )

    if isinstance(
        profile_data,
        dict,
    ):

        st.success(
            "Candidate profile is ready."
        )

        source = st.session_state.get(
            "candidate_source"
        )

        if source:
            st.caption(
                f"Source: {source}"
            )

        st.subheader(
            "Review Approved Candidate Profile"
        )

        full_name = st.text_input(
            "Full name",
            value=profile_data.get(
                "full_name",
                "",
            ),
            key="candidate_review_full_name",
        )

        professional_title = st.text_input(
            "Professional title",
            value=profile_data.get(
                "professional_title",
                "",
            ),
            key="candidate_review_title",
        )

        professional_summary = st.text_area(
            "Professional summary",
            value=profile_data.get(
                "professional_summary",
                "",
            ),
            height=150,
            key="candidate_review_summary",
        )

        technical_skills_text = st.text_area(
            "Technical skills — one per line",
            value="\n".join(
                profile_data.get(
                    "technical_skills",
                    [],
                )
            ),
            height=200,
            key="candidate_review_technical_skills",
        )

        programming_languages_text = st.text_area(
            "Programming languages — one per line",
            value="\n".join(
                profile_data.get(
                    "programming_languages",
                    [],
                )
            ),
            height=120,
            key="candidate_review_programming_languages",
        )

        tools_text = st.text_area(
            "Tools — one per line",
            value="\n".join(
                profile_data.get(
                    "tools",
                    [],
                )
            ),
            height=180,
            key="candidate_review_tools",
        )

        methodologies_text = st.text_area(
            "Methodologies — one per line",
            value="\n".join(
                profile_data.get(
                    "methodologies",
                    [],
                )
            ),
            height=120,
            key="candidate_review_methodologies",
        )

        industries_text = st.text_area(
            "Industries — one per line",
            value="\n".join(
                profile_data.get(
                    "industries",
                    [],
                )
            ),
            height=100,
            key="candidate_review_industries",
        )

        col_save_profile, col_download_profile = st.columns(
            2
        )

        with col_save_profile:

            if st.button(
                "Save Profile Changes",
                use_container_width=True,
            ):

                try:
                    updated_profile = (
                        profile_data.copy()
                    )

                    updated_profile[
                        "full_name"
                    ] = full_name

                    updated_profile[
                        "professional_title"
                    ] = professional_title

                    updated_profile[
                        "professional_summary"
                    ] = professional_summary

                    updated_profile[
                        "technical_skills"
                    ] = [
                        item.strip()
                        for item in technical_skills_text.splitlines()
                        if item.strip()
                    ]

                    updated_profile[
                        "programming_languages"
                    ] = [
                        item.strip()
                        for item in programming_languages_text.splitlines()
                        if item.strip()
                    ]

                    updated_profile[
                        "tools"
                    ] = [
                        item.strip()
                        for item in tools_text.splitlines()
                        if item.strip()
                    ]

                    updated_profile[
                        "methodologies"
                    ] = [
                        item.strip()
                        for item in methodologies_text.splitlines()
                        if item.strip()
                    ]

                    updated_profile[
                        "industries"
                    ] = [
                        item.strip()
                        for item in industries_text.splitlines()
                        if item.strip()
                    ]

                    updated_profile = (
                        validate_candidate_profile_data(
                            updated_profile
                        )
                    )

                    st.session_state[
                        "candidate_profile"
                    ] = updated_profile

                    clear_job_and_application_outputs()

                    save_candidate_profile_to_default_file(
                        updated_profile
                    )

                    st.success(
                        "Candidate profile changes saved."
                    )

                    st.rerun()

                except Exception as error:
                    show_error(
                        "Candidate profile changes could not be saved.",
                        error,
                    )

        with col_download_profile:

            profile_json = json.dumps(
                profile_data,
                indent=2,
                ensure_ascii=False,
            )

            st.download_button(
                "Download Candidate Profile",
                data=profile_json,
                file_name="candidate_profile.json",
                mime="application/json",
                use_container_width=True,
            )

        with st.expander(
            "View Complete Structured Profile"
        ):
            st.json(
                profile_data
            )

    else:
        st.info(
            "Load an existing candidate profile or analyse a CV to continue."
        )


# ============================================================
# JOB ANALYSIS TAB
# ============================================================

with job_tab:

    st.header(
        "Job Analysis"
    )

    candidate_data = (
        st.session_state.get(
            "candidate_profile"
        )
    )

    if not isinstance(
        candidate_data,
        dict,
    ):
        st.info(
            "Load or create a candidate profile in the Candidate Profile tab first."
        )

    else:

        st.success(
            "Candidate profile loaded."
        )

        job_description = st.text_area(
            "Job description",
            height=400,
            value=st.session_state.get(
                "job_description_text",
                "",
            ),
            placeholder=(
                "Paste the full job advertisement here..."
            ),
            key="job_description_input",
        )

        if st.button(
            "Analyse Job Description",
            type="primary",
        ):

            clean_job_description = (
                job_description.strip()
            )

            if not clean_job_description:
                st.warning(
                    "Please paste a job description first."
                )

            elif len(
                clean_job_description
            ) < 200:
                st.warning(
                    "The job description seems very short. "
                    "Please paste the complete vacancy."
                )

            else:

                with st.spinner(
                    "Career Copilot is analysing the job..."
                ):

                    try:
                        job_profile = (
                            create_job_profile(
                                clean_job_description
                            )
                        )

                        st.session_state[
                            "job_profile"
                        ] = (
                            job_profile.model_dump()
                        )

                        st.session_state[
                            "job_description_text"
                        ] = clean_job_description

                        clear_application_outputs()

                        # Calculate the deterministic C++ match immediately
                        # after the structured job profile is available.
                        match_result = (
                            calculate_cpp_match(
                                candidate_data,
                                st.session_state[
                                    "job_profile"
                                ],
                            )
                        )

                        st.session_state[
                            "match_result"
                        ] = match_result

                        st.success(
                            "Job description analysed successfully."
                        )

                        st.rerun()

                    except Exception as error:
                        show_error(
                            "Job analysis or compatibility calculation failed.",
                            error,
                        )

        job_data = (
            st.session_state.get(
                "job_profile"
            )
        )

        match_data = (
            st.session_state.get(
                "match_result"
            )
        )

        if isinstance(
            job_data,
            dict,
        ):

            st.divider()

            st.subheader(
                "Structured Job Profile"
            )

            job_col1, job_col2, job_col3 = st.columns(
                3
            )

            with job_col1:
                st.metric(
                    "Role",
                    job_data.get(
                        "job_title",
                        "",
                    )
                    or "Not specified",
                )

            with job_col2:
                st.metric(
                    "Company",
                    job_data.get(
                        "company",
                        "",
                    )
                    or "Not specified",
                )

            with job_col3:
                st.metric(
                    "Location",
                    job_data.get(
                        "location",
                        "",
                    )
                    or "Not specified",
                )

            with st.expander(
                "View Full Structured Job Profile"
            ):
                st.json(
                    job_data
                )

        if isinstance(
            match_data,
            dict,
        ):

            st.divider()

            st.subheader(
                "C++ Compatibility Analysis"
            )

            score = float(
                match_data.get(
                    "score",
                    0,
                )
            )

            matched_skills = (
                match_data.get(
                    "matched_skills",
                    [],
                )
            )

            missing_skills = (
                match_data.get(
                    "missing_skills",
                    [],
                )
            )

            metric_col1, metric_col2, metric_col3 = st.columns(
                3
            )

            with metric_col1:
                st.metric(
                    "Compatibility",
                    f"{score:.1f}%",
                )

            with metric_col2:
                st.metric(
                    "Matched Skills",
                    len(
                        matched_skills
                    ),
                )

            with metric_col3:
                st.metric(
                    "Missing Skills",
                    len(
                        missing_skills
                    ),
                )

            normalized_score = min(
                max(
                    score / 100,
                    0,
                ),
                1,
            )

            st.progress(
                normalized_score
            )

            skill_col1, skill_col2 = st.columns(
                2
            )

            with skill_col1:

                st.markdown(
                    "### Matched Skills"
                )

                if matched_skills:
                    for skill in matched_skills:
                        st.write(
                            f"✅ {skill}"
                        )

                else:
                    st.write(
                        "No exact skill matches found."
                    )

            with skill_col2:

                st.markdown(
                    "### Missing Skills"
                )

                if missing_skills:
                    for skill in missing_skills:
                        st.write(
                            f"⚠️ {skill}"
                        )

                else:
                    st.write(
                        "No missing required skills."
                    )


# ============================================================
# APPLICATION PACKAGE TAB
# ============================================================

with application_tab:

    st.header(
        "Application Package"
    )

    candidate_data = (
        st.session_state.get(
            "candidate_profile"
        )
    )

    job_data = (
        st.session_state.get(
            "job_profile"
        )
    )

    match_data = (
        st.session_state.get(
            "match_result"
        )
    )

    package_ready = (
        isinstance(
            candidate_data,
            dict,
        )
        and isinstance(
            job_data,
            dict,
        )
        and isinstance(
            match_data,
            dict,
        )
    )

    if not package_ready:

        st.info(
            "Complete Candidate Profile and Job Analysis first. "
            "The application package becomes available after the "
            "C++ compatibility analysis has completed."
        )

    else:

        st.success(
            "Candidate profile, job profile and compatibility result are ready."
        )

        target_role = (
            job_data.get(
                "job_title",
                "",
            )
            or "Current Job"
        )

        target_company = (
            job_data.get(
                "company",
                "",
            )
            or "Company not specified"
        )

        st.caption(
            f"Target: {target_role} — {target_company}"
        )

        cv_tab, cover_tab, interview_tab = st.tabs(
            [
                "Tailored CV",
                "Cover Letter",
                "Interview Prep",
            ]
        )

        # ----------------------------------------------------
        # TAILORED CV
        # ----------------------------------------------------

        with cv_tab:

            st.subheader(
                "Tailored CV"
            )

            if st.button(
                (
                    "Regenerate Tailored CV"
                    if isinstance(
                        st.session_state.get(
                            "tailored_cv"
                        ),
                        dict,
                    )
                    else "Generate Tailored CV"
                ),
                type="primary",
                key="generate_tailored_cv",
            ):

                with st.spinner(
                    "Career Copilot is tailoring your CV..."
                ):

                    try:
                        tailored_cv = (
                            create_tailored_cv(
                                candidate_data,
                                job_data,
                                match_data,
                            )
                        )

                        st.session_state[
                            "tailored_cv"
                        ] = (
                            tailored_cv.model_dump()
                        )

                        st.success(
                            "Tailored CV generated successfully."
                        )

                        st.rerun()

                    except Exception as error:
                        show_error(
                            "Tailored CV generation failed.",
                            error,
                        )

            cv_data = (
                st.session_state.get(
                    "tailored_cv"
                )
            )

            if isinstance(
                cv_data,
                dict,
            ):

                professional_title = st.text_input(
                    "Professional title",
                    value=cv_data.get(
                        "professional_title",
                        "",
                    ),
                    key="tailored_cv_title",
                )

                professional_summary = st.text_area(
                    "Professional summary",
                    value=cv_data.get(
                        "professional_summary",
                        "",
                    ),
                    height=180,
                    key="tailored_cv_summary",
                )

                prioritized_skills_text = st.text_area(
                    "Prioritized skills — one per line",
                    value="\n".join(
                        cv_data.get(
                            "prioritized_skills",
                            [],
                        )
                    ),
                    height=180,
                    key="tailored_cv_skills",
                )

                if st.button(
                    "Save CV Changes",
                    key="save_cv_changes",
                ):

                    updated_cv = (
                        cv_data.copy()
                    )

                    updated_cv[
                        "professional_title"
                    ] = professional_title

                    updated_cv[
                        "professional_summary"
                    ] = professional_summary

                    updated_cv[
                        "prioritized_skills"
                    ] = [
                        item.strip()
                        for item in prioritized_skills_text.splitlines()
                        if item.strip()
                    ]

                    st.session_state[
                        "tailored_cv"
                    ] = updated_cv

                    st.success(
                        "CV changes saved."
                    )

                    st.rerun()

                with st.expander(
                    "View Complete Tailored CV Structure"
                ):
                    st.json(
                        st.session_state[
                            "tailored_cv"
                        ]
                    )

                cv_docx = (
                    create_cv_docx(
                        st.session_state[
                            "tailored_cv"
                        ]
                    )
                )

                st.download_button(
                    "Download Tailored CV",
                    data=cv_docx,
                    file_name="tailored_cv.docx",
                    mime=(
                        "application/"
                        "vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                )

            else:
                st.info(
                    "Generate a tailored CV to view and download it."
                )

        # ----------------------------------------------------
        # COVER LETTER
        # ----------------------------------------------------

        with cover_tab:

            st.subheader(
                "Cover Letter"
            )

            if st.button(
                (
                    "Regenerate Cover Letter"
                    if isinstance(
                        st.session_state.get(
                            "cover_letter"
                        ),
                        str,
                    )
                    else "Generate Cover Letter"
                ),
                type="primary",
                key="generate_cover_letter",
            ):

                with st.spinner(
                    "Career Copilot is writing the cover letter..."
                ):

                    try:
                        cover_letter = (
                            create_cover_letter(
                                candidate_data,
                                job_data,
                                match_data,
                            )
                        )

                        st.session_state[
                            "cover_letter"
                        ] = cover_letter

                        st.success(
                            "Cover letter generated successfully."
                        )

                        st.rerun()

                    except Exception as error:
                        show_error(
                            "Cover letter generation failed.",
                            error,
                        )

            cover_letter_data = (
                st.session_state.get(
                    "cover_letter"
                )
            )

            if isinstance(
                cover_letter_data,
                str,
            ):

                edited_cover_letter = (
                    st.text_area(
                        "Cover Letter",
                        value=cover_letter_data,
                        height=600,
                        key="cover_letter_editor",
                    )
                )

                if st.button(
                    "Save Cover Letter Changes",
                    key="save_cover_letter_changes",
                ):

                    st.session_state[
                        "cover_letter"
                    ] = edited_cover_letter

                    st.success(
                        "Cover letter changes saved."
                    )

                    st.rerun()

                cover_docx = (
                    create_cover_letter_docx(
                        edited_cover_letter
                    )
                )

                st.download_button(
                    "Download Cover Letter",
                    data=cover_docx,
                    file_name="cover_letter.docx",
                    mime=(
                        "application/"
                        "vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                )

            else:
                st.info(
                    "Generate a cover letter to view and download it."
                )

        # ----------------------------------------------------
        # INTERVIEW PREPARATION
        # ----------------------------------------------------

        with interview_tab:

            st.subheader(
                "Interview Preparation"
            )

            if st.button(
                (
                    "Regenerate Interview Preparation"
                    if isinstance(
                        st.session_state.get(
                            "interview_preparation"
                        ),
                        dict,
                    )
                    else "Generate Interview Preparation"
                ),
                type="primary",
                key="generate_interview_preparation",
            ):

                with st.spinner(
                    "Career Copilot is preparing interview questions..."
                ):

                    try:
                        interview_prep = (
                            create_interview_preparation(
                                candidate_data,
                                job_data,
                                match_data,
                            )
                        )

                        st.session_state[
                            "interview_preparation"
                        ] = (
                            interview_prep.model_dump()
                        )

                        st.success(
                            "Interview preparation generated successfully."
                        )

                        st.rerun()

                    except Exception as error:
                        show_error(
                            "Interview preparation generation failed.",
                            error,
                        )

            prep_data = (
                st.session_state.get(
                    "interview_preparation"
                )
            )

            if isinstance(
                prep_data,
                dict,
            ):

                render_interview_questions(
                    prep_data
                )

                try:
                    interview_docx = (
                        create_interview_docx(
                            prep_data
                        )
                    )

                    st.download_button(
                        "Download Interview Preparation",
                        data=interview_docx,
                        file_name=(
                            "interview_preparation.docx"
                        ),
                        mime=(
                            "application/"
                            "vnd.openxmlformats-officedocument."
                            "wordprocessingml.document"
                        ),
                    )

                except Exception as error:
                    show_error(
                        "Interview preparation document creation failed.",
                        error,
                    )

            else:
                st.info(
                    "Generate interview preparation to view and download it."
                )

        # ----------------------------------------------------
        # SAVE APPLICATION
        # ----------------------------------------------------

        st.divider()

        st.subheader(
            "Save Application"
        )

        application_status = st.selectbox(
            "Application status",
            [
                "Analysed",
                "Prepared",
                "Applied",
                "Interview",
                "Rejected",
                "Offer",
                "Archived",
            ],
            index=1,
            key="application_status",
        )

        if st.button(
            "Save Application Record",
            key="save_application_record",
        ):

            try:
                saved_path = (
                    save_application_record(
                        candidate_data,
                        job_data,
                        match_data,
                        st.session_state.get(
                            "tailored_cv"
                        ),
                        st.session_state.get(
                            "cover_letter"
                        ),
                        st.session_state.get(
                            "interview_preparation"
                        ),
                        application_status,
                    )
                )

                st.success(
                    "Application saved successfully."
                )

                st.caption(
                    f"Saved to: {saved_path}"
                )

            except Exception as error:
                show_error(
                    "The application could not be saved.",
                    error,
                )


# ============================================================
# APPLICATION HISTORY TAB
# ============================================================

with history_tab:

    st.header(
        "Application History"
    )

    applications = (
        load_application_history()
    )

    if not applications:

        st.info(
            "No saved applications yet."
        )

    else:

        history_rows = []

        for application in applications:

            history_rows.append(
                {
                    "Company": application.get(
                        "company",
                        "",
                    ),
                    "Position": application.get(
                        "job_title",
                        "",
                    ),
                    "Location": application.get(
                        "location",
                        "",
                    ),
                    "Match Score": application.get(
                        "match_score",
                        0,
                    ),
                    "Status": application.get(
                        "status",
                        "",
                    ),
                    "Created": application.get(
                        "created_at",
                        "",
                    ),
                }
            )

        history_df = (
            pd.DataFrame(
                history_rows
            )
        )

        st.dataframe(
            history_df,
            use_container_width=True,
            hide_index=True,
        )
