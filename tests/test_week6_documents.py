import json
from pathlib import Path
from io import BytesIO

from docx import Document

from documents.application_docs import create_cv_docx, create_cover_letter_docx


PROFILE = json.loads(Path("candidate_profile.json").read_text(encoding="utf-8"))


def _cv_data():
    return {
        "professional_title": PROFILE["professional_title"],
        "professional_summary": PROFILE["professional_summary"],
        "prioritized_skills": ["C++", "AUTOSAR", "ISO 26262", "MATLAB/Simulink"],
        "experience": [
            {
                "company": entry["company"],
                "job_title": entry["role"],
                "bullets": entry["responsibilities"],
            }
            for entry in PROFILE["employment_history"]
        ],
        "education": [],
        "certifications": PROFILE["certifications"],
    }


def test_cv_docx_contains_reference_sections_and_identity():
    output = create_cv_docx(_cv_data(), PROFILE)
    assert output[:2] == b"PK"
    document = Document(BytesIO(output))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "TECHNICAL SKILLS" in text
    assert "WORK EXPERIENCE" in text
    assert "EDUCATION" in text
    assert "PATENTS & PUBLICATIONS & AWARDS" in text
    assert "LANGUAGES" in text
    assert PROFILE["email"] in text


def test_cover_letter_has_reference_structure():
    body = (
        "My systems engineering background aligns with this role.\n\n"
        "I can contribute embedded, simulation and validation experience.\n\n"
        "I would welcome the opportunity to discuss the position."
    )
    output = create_cover_letter_docx(
        body,
        PROFILE,
        {"company": "Example GmbH", "job_title": "Systems Engineer", "job_id": "123"},
    )
    document = Document(BytesIO(output))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "Example GmbH" in text
    assert "Subject: Application for Systems Engineer" in text
    assert "Dear Hiring Team," in text
    assert "Sincerely" in text
